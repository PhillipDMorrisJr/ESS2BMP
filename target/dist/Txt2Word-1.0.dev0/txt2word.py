import docx, sys

def txt2word(infile, outfile, tabStopWidth):
 headerOne = 1
 headerTwo = 2
 txtFile = open(infile)
 doc = docx.Document()
 text = txtFile.readlines() 
 prevline = ""
 currParagraph = ""

 for line in text:
 
  lastIndex = len(text) - 1
  currentIndex = text.index(line)
  
  if not lastIndex == currentIndex:
   nextIndex = text.index(line) + 1
   nextline = text[nextIndex]
   
   if meetsHeaderOneCondition(line, prevline, nextline):
    aline = line.replace("\t","")
    doc.add_paragraph(currParagraph)
    currParagraph = ""
    doc.add_heading(aline, headerOne)
    
   elif  meetsHeaderTwoCondition(line, prevline, nextline):
    doc.add_paragraph(currParagraph)
    currParagraph = ""
    doc.add_heading(line, headerTwo)
   elif line != "\n" and line != "\t":
    line.replace("\t", " "*tabStopWidth)
    if not line.endswith(".\n"):
     line.replace("\n", "")  
     currParagraph += line   
    else:
     currParagraph += line   
     doc.add_paragraph(currParagraph)
     currParagraph = ""
  else:
   currParagraph += line 
   doc.add_paragraph(currParagraph)
  prevline = line                                 

 doc.save(outfile)

def meetsHeaderOneCondition(line, prevline, nextline):
 return line.startswith("\t") and line.endswith("\n") and prevline == "\n" and nextline == "\n" and line != "\n" and line != "\t" 

def meetsHeaderTwoCondition(line, prevline, nextline):
 return prevline == "\n" and nextline == "\n" and not line.startswith("\t") and line != "\n" and line != "\t"

if __name__ == '__main__':
 # figure out what your command-line arguments are using sys.argv:
 infile = sys.argv[1]
 outfile = sys.argv[2]
 tabStopWidth = int(sys.argv[3])
 txt2word(infile, outfile, tabStopWidth)


